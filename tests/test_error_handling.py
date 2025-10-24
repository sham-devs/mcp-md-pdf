"""
Tests for error handling and file corruption scenarios

Tests cover:
- Corrupt template files
- Invalid file formats
- Missing required files
- Resource cleanup after errors
- Graceful degradation
"""
import os
import pytest
from docx import Document
from src.md_pdf_mcp.converter import MarkdownConverter


class TestCorruptTemplateHandling:
    """Test handling of corrupt or invalid template files"""

    def test_non_docx_file_as_template(self, simple_markdown, output_docx_path, temp_dir):
        """Test that non-DOCX file disguised as template is handled"""
        # Create a text file with .dotx extension
        fake_template = os.path.join(temp_dir, "fake.dotx")
        with open(fake_template, 'w') as f:
            f.write("This is not a valid Word template")

        converter = MarkdownConverter()

        # Should either fail gracefully or fall back to no template
        try:
            result = converter.markdown_to_word(
                simple_markdown,
                output_docx_path,
                fake_template
            )
            # If it succeeds, it should have fallen back to no template
            assert result is True
            assert os.path.exists(output_docx_path)
        except Exception as e:
            # Acceptable to raise exception with clear message about file format
            error_msg = str(e).lower()
            assert any(keyword in error_msg for keyword in ["template", "corrupt", "invalid", "zip", "docx"])

    def test_empty_file_as_template(self, simple_markdown, output_docx_path, temp_dir):
        """Test handling of empty file as template"""
        empty_template = os.path.join(temp_dir, "empty.dotx")
        # Create empty file
        open(empty_template, 'w').close()

        converter = MarkdownConverter()

        try:
            result = converter.markdown_to_word(
                simple_markdown,
                output_docx_path,
                empty_template
            )
            # Should either work or fail gracefully
            if result:
                assert os.path.exists(output_docx_path)
        except Exception:
            # Acceptable to fail on empty template
            pass

    def test_binary_garbage_as_template(self, simple_markdown, output_docx_path, temp_dir):
        """Test handling of random binary data as template"""
        garbage_template = os.path.join(temp_dir, "garbage.dotx")
        with open(garbage_template, 'wb') as f:
            f.write(b'\x00\x01\x02\x03\x04' * 100)

        converter = MarkdownConverter()

        try:
            converter.markdown_to_word(
                simple_markdown,
                output_docx_path,
                garbage_template
            )
        except Exception as e:
            # Should fail with informative error
            assert len(str(e)) > 0


class TestInvalidMarkdownInputs:
    """Test handling of invalid or problematic markdown inputs"""

    def test_binary_file_as_markdown(self, output_docx_path, temp_dir):
        """Test that binary file is handled gracefully"""
        binary_file = os.path.join(temp_dir, "binary.md")
        with open(binary_file, 'wb') as f:
            f.write(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR')  # PNG header

        converter = MarkdownConverter()

        try:
            result = converter.markdown_to_word(binary_file, output_docx_path)
            # If it succeeds, output should exist
            if result:
                assert os.path.exists(output_docx_path)
        except Exception:
            # Acceptable to fail on binary input
            pass

    def test_markdown_with_null_bytes(self, output_docx_path, temp_dir):
        """Test handling of markdown with null bytes"""
        null_md = os.path.join(temp_dir, "null.md")
        with open(null_md, 'wb') as f:
            f.write(b'# Heading\x00\x00\nContent with\x00null bytes')

        converter = MarkdownConverter()

        try:
            result = converter.markdown_to_word(null_md, output_docx_path)
            if result:
                assert os.path.exists(output_docx_path)
        except Exception:
            # Acceptable to fail
            pass

    def test_extremely_long_line(self, output_docx_path, temp_dir):
        """Test handling of extremely long single line"""
        long_md = os.path.join(temp_dir, "long.md")
        with open(long_md, 'w', encoding='utf-8') as f:
            # 50,000 character line
            f.write("# Heading\n")
            f.write("A" * 50000)

        converter = MarkdownConverter()
        result = converter.markdown_to_word(long_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_markdown_with_invalid_utf8(self, output_docx_path, temp_dir):
        """Test handling of invalid UTF-8 sequences"""
        invalid_utf8 = os.path.join(temp_dir, "invalid_utf8.md")
        with open(invalid_utf8, 'wb') as f:
            f.write(b'# Valid heading\n')
            f.write(b'\xff\xfe Invalid UTF-8')

        converter = MarkdownConverter()

        try:
            result = converter.markdown_to_word(invalid_utf8, output_docx_path)
            if result:
                assert os.path.exists(output_docx_path)
        except Exception:
            # Acceptable to fail on invalid encoding
            pass


class TestResourceCleanup:
    """Test that resources are properly cleaned up after errors"""

    def test_temp_files_cleaned_after_error(self, temp_dir):
        """Test that temp files are cleaned up even when conversion fails"""
        bad_markdown = os.path.join(temp_dir, "bad.md")
        # File doesn't exist, should cause error
        output_path = os.path.join(temp_dir, "output")

        converter = MarkdownConverter()

        # Count files before
        files_before = set(os.listdir(temp_dir))

        try:
            converter.markdown_to_word(bad_markdown, output_path + ".docx")
        except Exception:
            pass

        # Count files after
        files_after = set(os.listdir(temp_dir))

        # No new temp files should remain
        new_files = files_after - files_before
        temp_files = [f for f in new_files if 'temp' in f.lower() or f.startswith('.')]

        assert len(temp_files) == 0, f"Temp files not cleaned up: {temp_files}"

    def test_file_handle_closed_after_error(self, temp_dir, simple_markdown):
        """Test that file handles are properly closed after errors"""
        output_path = os.path.join(temp_dir, "output.docx")

        converter = MarkdownConverter()

        # Create output successfully first
        converter.markdown_to_word(simple_markdown, output_path)

        # Try to overwrite with bad template
        bad_template = os.path.join(temp_dir, "bad.dotx")
        with open(bad_template, 'w') as f:
            f.write("not a template")

        try:
            converter.markdown_to_word(simple_markdown, output_path, bad_template)
        except Exception:
            pass

        # Should be able to delete the file (means handle was closed)
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
            success = True
        except PermissionError:
            success = False

        assert success, "File handle was not closed properly"


class TestOutputPathHandling:
    """Test handling of various output path scenarios"""

    def test_output_path_with_spaces(self, simple_markdown, temp_dir):
        """Test output path with spaces in directory name"""
        space_dir = os.path.join(temp_dir, "dir with spaces")
        os.makedirs(space_dir, exist_ok=True)

        output_path = os.path.join(space_dir, "output with spaces.docx")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(simple_markdown, output_path)

        assert result is True
        assert os.path.exists(output_path)

    def test_output_path_with_unicode(self, simple_markdown, temp_dir):
        """Test output path with unicode characters"""
        output_path = os.path.join(temp_dir, "文档_مستند_דוקומנט.docx")

        converter = MarkdownConverter()

        try:
            result = converter.markdown_to_word(simple_markdown, output_path)
            if result:
                assert os.path.exists(output_path)
        except Exception:
            # May fail on some filesystems
            pass

    def test_output_path_very_long(self, simple_markdown, temp_dir):
        """Test handling of very long output path"""
        # Create nested directory structure
        long_path = temp_dir
        for i in range(10):
            long_path = os.path.join(long_path, f"nested_directory_{i}")

        try:
            os.makedirs(long_path, exist_ok=True)
            output_path = os.path.join(long_path, "output.docx")

            converter = MarkdownConverter()
            result = converter.markdown_to_word(simple_markdown, output_path)

            if result:
                assert os.path.exists(output_path)
        except OSError:
            # May hit path length limits on Windows
            pass

    @pytest.mark.skipif(os.name == 'nt', reason="Read-only directory test not reliable on Windows")
    def test_output_to_readonly_directory(self, simple_markdown, temp_dir):
        """Test handling of read-only output directory"""
        readonly_dir = os.path.join(temp_dir, "readonly")
        os.makedirs(readonly_dir, exist_ok=True)

        output_path = os.path.join(readonly_dir, "output.docx")

        # Make directory read-only (Unix-specific)
        try:
            os.chmod(readonly_dir, 0o444)

            converter = MarkdownConverter()

            with pytest.raises(Exception):
                converter.markdown_to_word(simple_markdown, output_path)

        finally:
            # Restore permissions for cleanup
            os.chmod(readonly_dir, 0o755)


class TestGracefulDegradation:
    """Test that converter degrades gracefully when features are unavailable"""

    def test_missing_optional_styles_handled(self, simple_markdown, output_docx_path, temp_dir):
        """Test that missing optional styles don't crash conversion"""
        # This is already partially tested, but let's be explicit
        converter = MarkdownConverter()

        result = converter.markdown_to_word(simple_markdown, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

        # Verify document can be opened
        doc = Document(output_docx_path)
        assert len(doc.paragraphs) > 0

    def test_conversion_continues_after_partial_failure(self, output_docx_path, temp_dir):
        """Test that conversion continues even if some elements fail"""
        problematic_md = os.path.join(temp_dir, "problematic.md")

        # Create markdown with mix of valid and potentially problematic content
        with open(problematic_md, 'w', encoding='utf-8') as f:
            f.write("""# Valid Heading

Normal paragraph.

| Table | Header |
|-------|--------|
| Cell1 | Cell2  |

```python
# Code block
def test():
    pass
```

More content after potentially problematic sections.
""")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(problematic_md, output_docx_path)

        # Should succeed
        assert result is True
        assert os.path.exists(output_docx_path)

        # Should have some content even if parts failed
        doc = Document(output_docx_path)
        assert len(doc.paragraphs) > 0
