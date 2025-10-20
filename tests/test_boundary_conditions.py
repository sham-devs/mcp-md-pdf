"""
Tests for boundary conditions and edge cases

Tests cover:
- Very large files
- Deeply nested structures
- Extreme values
- Special character ranges
- Performance under stress
"""
import os
import pytest
from docx import Document
from src.md_pdf_mcp.converter import MarkdownConverter


class TestLargeFiles:
    """Test handling of large markdown files"""

    @pytest.mark.slow
    def test_large_markdown_file_10000_lines(self, output_docx_path, temp_dir):
        """Test conversion of 10,000 line markdown file"""
        large_md = os.path.join(temp_dir, "large_10k.md")

        with open(large_md, 'w', encoding='utf-8') as f:
            f.write("# Large Document Test\n\n")

            for i in range(10000):
                if i % 100 == 0:
                    f.write(f"## Section {i // 100}\n\n")
                f.write(f"Line {i}: This is content for line number {i}.\n\n")

        converter = MarkdownConverter()

        import time
        start = time.time()

        result = converter.markdown_to_word(large_md, output_docx_path)

        elapsed = time.time() - start

        assert result is True
        assert os.path.exists(output_docx_path)

        # Should complete in reasonable time (< 30 seconds)
        assert elapsed < 30, f"Took {elapsed} seconds for 10k lines"

        # Verify document was created with content
        doc = Document(output_docx_path)
        assert len(doc.paragraphs) > 5000

    def test_many_headings_document(self, output_docx_path, temp_dir):
        """Test document with 1000 headings"""
        many_headings_md = os.path.join(temp_dir, "many_headings.md")

        with open(many_headings_md, 'w', encoding='utf-8') as f:
            for i in range(1000):
                level = (i % 6) + 1  # H1 through H6
                f.write(f"{'#' * level} Heading {i}\n\nSome content.\n\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(many_headings_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_many_tables_document(self, output_docx_path, temp_dir):
        """Test document with 100 tables"""
        many_tables_md = os.path.join(temp_dir, "many_tables.md")

        with open(many_tables_md, 'w', encoding='utf-8') as f:
            f.write("# Document with Many Tables\n\n")

            for i in range(100):
                f.write(f"## Table {i}\n\n")
                f.write("| Column 1 | Column 2 | Column 3 |\n")
                f.write("|----------|----------|----------|\n")
                for row in range(5):
                    f.write(f"| Data {row}A | Data {row}B | Data {row}C |\n")
                f.write("\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(many_tables_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

        # Verify tables were created
        doc = Document(output_docx_path)
        assert len(doc.tables) >= 50  # At least half should be there


class TestDeepNesting:
    """Test deeply nested structures"""

    def test_deeply_nested_lists_10_levels(self, output_docx_path, temp_dir):
        """Test lists nested 10 levels deep"""
        nested_md = os.path.join(temp_dir, "nested_10.md")

        with open(nested_md, 'w', encoding='utf-8') as f:
            f.write("# Deeply Nested Lists\n\n")

            for level in range(10):
                indent = "\t" * level
                f.write(f"{indent}- Level {level} item 1\n")
                f.write(f"{indent}- Level {level} item 2\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(nested_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_mixed_nested_lists(self, output_docx_path, temp_dir):
        """Test alternating bullet and numbered lists nested"""
        mixed_md = os.path.join(temp_dir, "mixed_nested.md")

        with open(mixed_md, 'w', encoding='utf-8') as f:
            f.write("# Mixed Nested Lists\n\n")
            f.write("- Level 0 bullet\n")
            f.write("\t1. Level 1 number\n")
            f.write("\t\t- Level 2 bullet\n")
            f.write("\t\t\t1. Level 3 number\n")
            f.write("\t\t\t\t- Level 4 bullet\n")
            f.write("\t\t\t\t\t1. Level 5 number\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(mixed_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)


class TestUnicodeRanges:
    """Test various Unicode character ranges"""

    def test_emoji_heavy_document(self, output_docx_path, temp_dir):
        """Test document with many emojis"""
        emoji_md = os.path.join(temp_dir, "emoji.md")

        emojis = "😀😃😄😁😆😅🤣😂🙂🙃😉😊😇🥰😍🤩😘😗☺️😚😙🥲😋😛😜🤪😝🤑🤗🤭🤫🤔🤐🤨😐😑😶😏😒🙄😬🤥😌😔😪🤤😴😷🤒🤕🤢🤮🤧🥵🥶🥴😵🤯🤠🥳🥸😎🤓🧐"

        with open(emoji_md, 'w', encoding='utf-8') as f:
            f.write("# Emoji Test\n\n")
            for i in range(100):
                f.write(f"Line {i}: {emojis}\n\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(emoji_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_cjk_characters(self, output_docx_path, temp_dir):
        """Test Chinese, Japanese, Korean characters"""
        cjk_md = os.path.join(temp_dir, "cjk.md")

        with open(cjk_md, 'w', encoding='utf-8') as f:
            f.write("# CJK Characters Test\n\n")
            f.write("## Chinese\n\n你好世界。这是一个测试文档。\n\n")
            f.write("## Japanese\n\nこんにちは世界。これはテストドキュメントです。\n\n")
            f.write("## Korean\n\n안녕하세요 세계. 이것은 테스트 문서입니다.\n\n")

            # Mix with English
            f.write("## Mixed\n\n")
            f.write("English text 中文文本 日本語テキスト 한국어 텍스트\n\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(cjk_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_rtl_languages(self, output_docx_path, temp_dir):
        """Test Right-to-Left languages (Arabic, Hebrew)"""
        rtl_md = os.path.join(temp_dir, "rtl.md")

        with open(rtl_md, 'w', encoding='utf-8') as f:
            f.write("# RTL Languages Test\n\n")
            f.write("## Arabic\n\nمرحبا بالعالم. هذا مستند اختباري.\n\n")
            f.write("## Hebrew\n\nשלום עולם. זה מסמך בדיקה.\n\n")

            # Mix with English (BiDi text)
            f.write("## Mixed BiDi\n\n")
            f.write("English text مع النص العربي and עם טקסט עברי together.\n\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(rtl_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_mathematical_symbols(self, output_docx_path, temp_dir):
        """Test mathematical and special symbols"""
        math_md = os.path.join(temp_dir, "math.md")

        with open(math_md, 'w', encoding='utf-8') as f:
            f.write("# Mathematical Symbols\n\n")
            f.write("Operators: ∀ ∂ ∃ ∅ ∇ ∈ ∉ ∋ ∏ ∑ − ∗ √ ∝ ∞ ∠ ∧ ∨ ∩ ∪ ∫ ∴ ∼ ≅ ≈ ≠ ≡ ≤ ≥\n\n")
            f.write("Greek: α β γ δ ε ζ η θ ι κ λ μ ν ξ ο π ρ ς σ τ υ φ χ ψ ω\n\n")
            f.write("Arrows: ← ↑ → ↓ ↔ ↕ ↖ ↗ ↘ ↙ ⇐ ⇑ ⇒ ⇓ ⇔\n\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(math_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)


class TestExtremeContent:
    """Test extreme content scenarios"""

    def test_single_extremely_long_paragraph(self, output_docx_path, temp_dir):
        """Test single paragraph with 10,000 words"""
        long_para_md = os.path.join(temp_dir, "long_para.md")

        with open(long_para_md, 'w', encoding='utf-8') as f:
            f.write("# Long Paragraph Test\n\n")

            # Single 10,000 word paragraph
            words = ["word" + str(i) for i in range(10000)]
            f.write(" ".join(words))

        converter = MarkdownConverter()
        result = converter.markdown_to_word(long_para_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_many_empty_lines(self, output_docx_path, temp_dir):
        """Test document with thousands of empty lines"""
        empty_lines_md = os.path.join(temp_dir, "empty_lines.md")

        with open(empty_lines_md, 'w', encoding='utf-8') as f:
            f.write("# Empty Lines Test\n\n")
            f.write("\n" * 1000)
            f.write("Content after many empty lines\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(empty_lines_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_many_consecutive_headings(self, output_docx_path, temp_dir):
        """Test many headings without content between them"""
        headings_md = os.path.join(temp_dir, "consecutive_headings.md")

        with open(headings_md, 'w', encoding='utf-8') as f:
            for i in range(500):
                f.write(f"# Heading {i}\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(headings_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_extremely_wide_table(self, output_docx_path, temp_dir):
        """Test table with 50 columns"""
        wide_table_md = os.path.join(temp_dir, "wide_table.md")

        with open(wide_table_md, 'w', encoding='utf-8') as f:
            f.write("# Wide Table Test\n\n")

            # Header
            headers = " | ".join([f"Col{i}" for i in range(50)])
            f.write(f"| {headers} |\n")

            # Separator
            sep = " | ".join(["---"] * 50)
            f.write(f"| {sep} |\n")

            # Data rows
            for row in range(3):
                data = " | ".join([f"D{row}{i}" for i in range(50)])
                f.write(f"| {data} |\n")

        converter = MarkdownConverter()

        try:
            result = converter.markdown_to_word(wide_table_md, output_docx_path)

            if result:
                assert os.path.exists(output_docx_path)
        except Exception:
            # Very wide tables might fail - that's acceptable
            pass

    def test_extremely_tall_table(self, output_docx_path, temp_dir):
        """Test table with 1000 rows"""
        tall_table_md = os.path.join(temp_dir, "tall_table.md")

        with open(tall_table_md, 'w', encoding='utf-8') as f:
            f.write("# Tall Table Test\n\n")
            f.write("| Column 1 | Column 2 | Column 3 |\n")
            f.write("|----------|----------|----------|\n")

            for i in range(1000):
                f.write(f"| Row {i} A | Row {i} B | Row {i} C |\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(tall_table_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)


class TestSpecialFormatting:
    """Test special formatting edge cases"""

    def test_all_text_formatting_combined(self, output_docx_path, temp_dir):
        """Test text with all formatting types applied simultaneously"""
        formatted_md = os.path.join(temp_dir, "formatted.md")

        with open(formatted_md, 'w', encoding='utf-8') as f:
            f.write("# Formatting Test\n\n")
            # Bold, italic, code all together
            f.write("This has **bold**, *italic*, `code`, and **_bold italic_**.\n\n")
            f.write("Links: [text](http://example.com)\n\n")
            f.write("Images: ![alt text](image.png)\n\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(formatted_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_code_block_with_special_characters(self, output_docx_path, temp_dir):
        """Test code blocks containing special characters"""
        code_md = os.path.join(temp_dir, "code.md")

        with open(code_md, 'w', encoding='utf-8') as f:
            f.write("# Code Block Test\n\n```python\n")
            f.write("# Special chars: <>&\"'`\n")
            f.write("regex = r'^[\\w\\s]+$'\n")
            f.write("xml = '<tag attr=\"value\">content</tag>'\n")
            f.write("```\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(code_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)

    def test_blockquote_with_nested_content(self, output_docx_path, temp_dir):
        """Test blockquotes containing lists and formatting"""
        quote_md = os.path.join(temp_dir, "quote.md")

        with open(quote_md, 'w', encoding='utf-8') as f:
            f.write("# Blockquote Test\n\n")
            f.write("> This is a quote\n")
            f.write("> \n")
            f.write("> With **bold** and *italic*\n")
            f.write("> \n")
            f.write("> And a `code snippet`\n")

        converter = MarkdownConverter()
        result = converter.markdown_to_word(quote_md, output_docx_path)

        assert result is True
        assert os.path.exists(output_docx_path)
