#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word/PDF Converter
Extracted and refactored from MedConnect documentation generator
"""

import os
import re
from pathlib import Path
from typing import Optional, List

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class MarkdownConverter:
    """Converts markdown files to Word/PDF with optional template support"""

    def __init__(self):
        self.image_cache = {}

    def _add_paragraph_shading(self, paragraph, fill_color):
        """Add background shading to a paragraph"""
        try:
            pPr = paragraph._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill_color)
            pPr.append(shd)
        except Exception:
            pass  # Fail gracefully if shading can't be applied

    def _add_paragraph_border(self, paragraph, border_positions, size='6', color='CCCCCC', space='4'):
        """Add borders to a paragraph

        Args:
            paragraph: The paragraph to add borders to
            border_positions: List of positions ('top', 'left', 'bottom', 'right', 'all')
            size: Border size in eighths of a point (default '6' = 0.75pt)
            color: Border color in hex (default 'CCCCCC')
            space: Space between border and text in points (default '4')
        """
        try:
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            positions = ['top', 'left', 'bottom', 'right'] if 'all' in border_positions else border_positions

            for pos in positions:
                border = OxmlElement(f'w:{pos}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), size)
                border.set(qn('w:space'), space)
                border.set(qn('w:color'), color)
                pBdr.append(border)

            pPr.append(pBdr)
        except Exception:
            pass  # Fail gracefully if borders can't be applied

    def _add_run_shading(self, run, fill_color):
        """Add background shading to a run (inline text)"""
        try:
            rPr = run._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill_color)
            rPr.append(shd)
        except Exception:
            pass  # Fail gracefully if shading can't be applied

    def markdown_to_word(
        self,
        markdown_path: str,
        output_path: str,
        template_path: Optional[str] = None
    ) -> bool:
        """
        Convert markdown file to Word document.

        Args:
            markdown_path: Path to input .md file
            output_path: Path for output .docx file
            template_path: Optional path to .dotx template

        Returns:
            True if successful, False otherwise
        """
        try:
            # Load template or create new document
            if template_path and os.path.exists(template_path):
                doc = self._load_template(template_path)
                # Create bullet numbering even when using template
                self._create_bullet_numbering(doc)
            else:
                doc = Document()
                self._apply_default_styles(doc)

            # Parse and add markdown content
            self._add_markdown_to_doc(doc, markdown_path)

            # Save document
            doc.save(output_path)
            return True

        except Exception as e:
            raise Exception(f"Failed to convert markdown to Word: {e}")

    def word_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """
        Convert Word document to PDF.

        Platform-specific implementations (priority order):
        1. Pandoc (if available) - lightweight and cross-platform
        2. Windows: Microsoft Word via COM automation
        3. Linux/macOS: LibreOffice in headless mode

        Args:
            docx_path: Path to input .docx file
            pdf_path: Path for output .pdf file

        Returns:
            True if successful, False otherwise
        """
        import platform
        import subprocess
        import shutil

        # Convert to absolute paths
        docx_path = os.path.abspath(docx_path)
        pdf_path = os.path.abspath(pdf_path)

        # DOCX→PDF conversion strategy: Use LibreOffice ONLY
        # LibreOffice preserves ALL DOCX formatting (colors, backgrounds, borders, tables)
        # Pandoc is NOT used for DOCX→PDF as it loses visual styling

        if platform.system() == "Windows":
            # Windows: Try Microsoft Word COM first (best formatting)
            if shutil.which('WINWORD.EXE') or self._is_word_installed():
                return self._word_to_pdf_windows(docx_path, pdf_path)

            # Fall back to LibreOffice if Word unavailable
            if shutil.which('soffice') or shutil.which('libreoffice'):
                return self._word_to_pdf_libreoffice(docx_path, pdf_path)
        else:
            # Linux/macOS: Use LibreOffice (preserves formatting perfectly)
            if shutil.which('soffice') or shutil.which('libreoffice'):
                return self._word_to_pdf_libreoffice(docx_path, pdf_path)

        # No suitable conversion tool available
        raise Exception(
            "LibreOffice required for DOCX→PDF conversion with formatting preservation.\n"
            "Install LibreOffice:\n"
            "  - Ubuntu/Debian: sudo apt-get install libreoffice\n"
            "  - macOS: brew install --cask libreoffice\n"
            "  - Windows: Download from https://www.libreoffice.org/download/"
        )

    def _is_word_installed(self) -> bool:
        """Check if Microsoft Word is installed on Windows"""
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Quit()
            return True
        except Exception:
            return False

    def _word_to_pdf_windows(self, docx_path: str, pdf_path: str) -> bool:
        """Convert Word to PDF using Microsoft Word COM automation (Windows only)"""
        try:
            import win32com.client

            # Create Word application
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            try:
                # Open document
                doc = word.Documents.Open(docx_path)

                # Update fields (TOC, etc.)
                doc.Fields.Update()

                # Save as PDF (wdFormatPDF = 17)
                doc.SaveAs(pdf_path, FileFormat=17)

                # Close document
                doc.Close(SaveChanges=False)

                return True

            finally:
                word.Quit()

        except ImportError:
            raise Exception("pywin32 library required for PDF conversion on Windows. Install: pip install pywin32")
        except Exception as e:
            raise Exception(f"Failed to convert Word to PDF using Microsoft Word: {e}")

    def _word_to_pdf_libreoffice(self, docx_path: str, pdf_path: str) -> bool:
        """Convert Word to PDF using LibreOffice (Linux/macOS)"""
        import subprocess
        import shutil

        # Check if LibreOffice is installed
        libreoffice_cmd = shutil.which('libreoffice') or shutil.which('soffice')

        if not libreoffice_cmd:
            raise Exception(
                "LibreOffice not found. Install LibreOffice for PDF conversion:\n"
                "  - Ubuntu/Debian: sudo apt-get install libreoffice\n"
                "  - macOS: brew install --cask libreoffice\n"
                "  - Fedora: sudo dnf install libreoffice"
            )

        try:
            # Get output directory
            output_dir = os.path.dirname(pdf_path)
            if not output_dir:
                output_dir = "."

            # Run LibreOffice headless conversion
            result = subprocess.run(
                [
                    libreoffice_cmd,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    docx_path
                ],
                capture_output=True,
                text=True,
                timeout=60  # 60 second timeout
            )

            # Check if conversion succeeded
            if result.returncode != 0:
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")

            # LibreOffice creates PDF with same basename as input
            expected_pdf = os.path.join(output_dir, os.path.basename(docx_path).replace('.docx', '.pdf'))

            # If output filename is different, rename it
            if expected_pdf != pdf_path and os.path.exists(expected_pdf):
                os.rename(expected_pdf, pdf_path)

            if not os.path.exists(pdf_path):
                raise Exception("PDF file was not created by LibreOffice")

            return True

        except subprocess.TimeoutExpired:
            raise Exception("LibreOffice conversion timed out (>60 seconds)")
        except Exception as e:
            raise Exception(f"Failed to convert Word to PDF using LibreOffice: {e}")

    def _word_to_pdf_pandoc(self, docx_path: str, pdf_path: str) -> bool:
        """Convert Word to PDF using Pandoc (cross-platform, lightweight)"""
        import subprocess

        try:
            # Run Pandoc conversion
            # Use XeLaTeX for full Unicode/emoji support
            result = subprocess.run(
                [
                    'pandoc',
                    docx_path,
                    '-o', pdf_path,
                    '--pdf-engine=xelatex'  # XeLaTeX has native Unicode support
                ],
                capture_output=True,
                text=True,
                timeout=60  # 60 second timeout
            )

            # Check if conversion succeeded
            if result.returncode != 0:
                raise Exception(f"Pandoc conversion failed: {result.stderr}")

            if not os.path.exists(pdf_path):
                raise Exception("PDF file was not created by Pandoc")

            return True

        except subprocess.TimeoutExpired:
            raise Exception("Pandoc conversion timed out (>60 seconds)")
        except FileNotFoundError:
            raise Exception(
                "Pandoc not found. Install Pandoc for PDF conversion:\n"
                "  - Ubuntu/Debian: sudo apt-get install pandoc texlive-latex-base\n"
                "  - macOS: brew install pandoc basictex\n"
                "  - Windows: choco install pandoc miktex\n"
                "  - Or download from: https://pandoc.org/installing.html"
            )
        except Exception as e:
            raise Exception(f"Failed to convert Word to PDF using Pandoc: {e}")

    def _load_template(self, template_path: str) -> Document:
        """Load .dotx template file"""
        import tempfile
        import shutil
        from zipfile import ZipFile
        import zipfile

        if not str(template_path).endswith('.dotx'):
            return Document(template_path)

        # Workaround for .dotx files
        temp_dir = tempfile.gettempdir()
        temp_template = os.path.join(temp_dir, 'template_temp.docx')

        # Copy .dotx to temp .docx
        shutil.copy2(str(template_path), temp_template)

        # Modify internal content type
        with ZipFile(temp_template, 'r') as zip_read:
            content_types_xml = zip_read.read('[Content_Types].xml')

        content_types_str = content_types_xml.decode('utf-8')
        content_types_str = content_types_str.replace(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
        )

        temp_modified = os.path.join(temp_dir, 'template_modified.docx')
        with ZipFile(temp_template, 'r') as zip_read:
            with ZipFile(temp_modified, 'w', zipfile.ZIP_DEFLATED) as zip_write:
                for item in zip_read.infolist():
                    if item.filename == '[Content_Types].xml':
                        zip_write.writestr(item, content_types_str.encode('utf-8'))
                    else:
                        zip_write.writestr(item, zip_read.read(item.filename))

        doc = Document(temp_modified)

        # Clean up temp files
        try:
            os.remove(temp_template)
            os.remove(temp_modified)
        except:
            pass

        return doc

    def _apply_default_styles(self, doc: Document):
        """Apply default styling to document"""
        # Create bullet numbering
        self._create_bullet_numbering(doc)

    def _create_bullet_numbering(self, doc: Document):
        """Create bullet list numbering definition"""
        from docx.oxml import parse_xml

        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            from docx.opc.part import XmlPart
            from docx.opc.constants import CONTENT_TYPE as CT
            numbering_part = doc.part.relate_to(
                XmlPart(CT.WML_NUMBERING),
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'
            )

        bullet_abstract_num_xml = '''
        <w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="100">
            <w:multiLevelType w:val="hybridMultilevel"/>
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="bullet"/>
                <w:lvlText w:val="●"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="357" w:hanging="357"/>
                    <w:spacing w:after="120"/>
                </w:pPr>
            </w:lvl>
            <w:lvl w:ilvl="1">
                <w:start w:val="1"/>
                <w:numFmt w:val="bullet"/>
                <w:lvlText w:val="●"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="714" w:hanging="357"/>
                    <w:spacing w:after="120"/>
                </w:pPr>
            </w:lvl>
            <w:lvl w:ilvl="2">
                <w:start w:val="1"/>
                <w:numFmt w:val="bullet"/>
                <w:lvlText w:val="○"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="1071" w:hanging="357"/>
                    <w:spacing w:after="120"/>
                </w:pPr>
            </w:lvl>
        </w:abstractNum>
        '''

        num_xml = '''
        <w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="100">
            <w:abstractNumId w:val="100"/>
        </w:num>
        '''

        try:
            numbering_element = numbering_part.element

            existing_abstract = numbering_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId="100"]')
            if existing_abstract is None:
                abstract_num = parse_xml(bullet_abstract_num_xml)
                numbering_element.insert(0, abstract_num)

            existing_num = numbering_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId="100"]')
            if existing_num is None:
                num = parse_xml(num_xml)
                numbering_element.append(num)

        except Exception as e:
            pass  # Silently fail if numbering creation fails

    def _parse_markdown_table(self, lines: List[str], start_index: int):
        """
        Parse a markdown table starting at start_index
        Returns: (table_rows, next_index) where table_rows is list of lists, next_index is line after table
        """
        table_rows = []
        i = start_index

        while i < len(lines):
            line = lines[i].strip()

            # Check if this is a table row
            if line.startswith('|') and line.endswith('|'):
                # Skip separator rows (|---|---|)
                if not line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '').strip():
                    i += 1
                    continue

                # Split by | and clean up cells
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            else:
                # End of table
                break

        return table_rows, i

    def _add_word_table(self, doc: Document, table_rows: List[List[str]]):
        """Create a properly formatted Word table from parsed markdown table rows"""
        if not table_rows or len(table_rows) < 1:
            return

        # Create table
        num_cols = len(table_rows[0])

        # Normalize all rows to have same number of columns
        # Pad short rows with empty strings, truncate long rows
        normalized_rows = []
        for row in table_rows:
            if len(row) < num_cols:
                # Pad with empty strings
                normalized_rows.append(row + [''] * (num_cols - len(row)))
            elif len(row) > num_cols:
                # Truncate to num_cols
                normalized_rows.append(row[:num_cols])
            else:
                normalized_rows.append(row)

        table = doc.add_table(rows=len(normalized_rows), cols=num_cols)

        # Apply basic table style (Table Grid is built-in)
        try:
            table.style = 'Table Grid'
        except:
            pass  # If style not available, use default

        # Fill table cells
        for row_idx, row_data in enumerate(normalized_rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)

                # Clear default paragraph and add formatted text
                paragraph = cell.paragraphs[0]
                self._add_formatted_text(paragraph, cell_text)

                # Format header row (first row) - make bold if not already formatted
                if row_idx == 0:
                    # If there are no runs (no formatting), make entire cell bold
                    if len(paragraph.runs) == 0:
                        run = paragraph.add_run(cell_text)
                        run.font.bold = True
                        run.font.size = Pt(11)
                    else:
                        # Apply bold to all runs in header
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                    # Add shading to header row
                    try:
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'D9E2F3')  # Light blue background
                        cell._element.get_or_add_tcPr().append(shading)
                    except:
                        pass  # If shading fails, continue without it

        # Set header row to repeat on each page
        if len(table.rows) > 0:
            try:
                table.rows[0]._tr.get_or_add_trPr().append(
                    OxmlElement('w:tblHeader')
                )
            except:
                pass

    def _restart_numbering(self, paragraph, doc: Document):
        """
        Restart numbering for a numbered list paragraph by creating a new numId
        with startOverride. Based on python-docx best practices.
        """
        try:
            # Get the numId that the paragraph's style uses
            styles = doc.styles
            style = paragraph.style

            # Get numId from the style
            num_id_current = -1
            try:
                if style._element.pPr is not None and style._element.pPr.numPr is not None:
                    num_id_current = style._element.pPr.numPr.numId.val
            except:
                # If style doesn't have numbering, try to get from paragraph
                if paragraph._element.pPr is not None and paragraph._element.pPr.numPr is not None:
                    num_id_current = paragraph._element.pPr.numPr.numId.val

            if num_id_current <= 0:
                return  # No numbering found

            # Get the numbering definitions
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return

            ct_numbering = numbering_part.numbering_definitions._numbering
            ct_num = ct_numbering.num_having_numId(num_id_current)

            if ct_num is None:
                return

            # Get the abstractNumId
            abstractNumId = ct_num.abstractNumId.val

            # Create a new num element with the same abstractNumId
            ct_num_new = ct_numbering.add_num(abstractNumId)
            num_id_new = ct_num_new.numId

            # Add startOverride to restart at 1
            lvl_override = ct_num_new.add_lvlOverride(0)  # Level 0
            start_override = lvl_override._add_startOverride()
            start_override.val = 1

            # Link the paragraph to the new numId
            numPr = paragraph._element.pPr._add_numPr()
            numPr._add_numId().val = num_id_new
            numPr._add_ilvl().val = 0  # Level 0

        except Exception as e:
            # Silently fail if numbering restart doesn't work
            pass

    def _add_markdown_to_doc(self, doc: Document, markdown_path: str):
        """Parse markdown file and add content to document"""
        with open(markdown_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        in_code_block = False
        code_content = []
        prev_style_type = None
        should_restart_numbering = False

        i = 0
        while i < len(lines):
            line = lines[i]
            style_type, content, level = self._parse_markdown_line(line)

            # Handle tables (must parse all rows together)
            if style_type == 'table':
                table_rows, next_i = self._parse_markdown_table(lines, i)
                if table_rows:
                    self._add_word_table(doc, table_rows)
                prev_style_type = style_type
                i = next_i
                continue

            # Detect when to restart numbering
            if style_type == 'numbered_list' and level == 0:
                # Restart if previous was not a numbered list or empty
                if prev_style_type not in ['numbered_list', 'empty', None]:
                    should_restart_numbering = True

            # Handle code blocks
            if style_type == 'code_marker':
                if in_code_block:
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph(code_text)
                    p.style = 'No Spacing'
                    p.paragraph_format.left_indent = Inches(0.5)

                    # Apply monospace font
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    # Add professional styling
                    self._add_paragraph_shading(p, 'F5F5F5')  # Light gray background
                    self._add_paragraph_border(p, ['all'], size='6', color='CCCCCC', space='4')

                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line.rstrip())
                i += 1
                continue

            # Add content based on style
            if style_type == 'Heading 1':
                doc.add_heading(content, level=1)
                prev_style_type = style_type
            elif style_type == 'Heading 2':
                doc.add_heading(content, level=2)
                prev_style_type = style_type
            elif style_type == 'Heading 3':
                doc.add_heading(content, level=3)
                prev_style_type = style_type
            elif style_type == 'Heading 4':
                try:
                    doc.add_heading(content, level=4)
                except KeyError:
                    # Style doesn't exist, use formatted paragraph instead
                    p = doc.add_paragraph()
                    run = p.add_run(content)
                    run.bold = True
                    run.font.size = Pt(12)
                prev_style_type = style_type
            elif style_type == 'Heading 5':
                try:
                    doc.add_heading(content, level=5)
                except KeyError:
                    # Style doesn't exist, use formatted paragraph instead
                    p = doc.add_paragraph()
                    run = p.add_run(content)
                    run.bold = True
                    run.font.size = Pt(11)
                prev_style_type = style_type
            elif style_type == 'Heading 6':
                try:
                    doc.add_heading(content, level=6)
                except KeyError:
                    # Style doesn't exist, use formatted paragraph instead
                    p = doc.add_paragraph()
                    run = p.add_run(content)
                    run.bold = True
                    run.font.size = Pt(10)
                prev_style_type = style_type
            elif style_type == 'hr':
                # Skip horizontal rules - not needed in Word document
                prev_style_type = style_type
            elif style_type == 'numbered_list':
                # Try to use list style, fall back to Normal if not available
                style_to_use = 'List Number' if level == 0 else f'List Number {level + 1}'
                try:
                    p = doc.add_paragraph(style=style_to_use)
                except KeyError:
                    # Style doesn't exist in template, use Normal style
                    p = doc.add_paragraph(style='Normal')

                # Restart numbering if needed
                if should_restart_numbering and level == 0:
                    self._restart_numbering(p, doc)
                    should_restart_numbering = False

                self._add_formatted_text(p, content)
                prev_style_type = style_type
            elif style_type == 'list':
                p = doc.add_paragraph(style='Normal')

                # Add bullet numbering
                pPr = p._element.get_or_add_pPr()
                numPr = OxmlElement('w:numPr')
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), '100')
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'), str(level))
                numPr.append(ilvl)
                numPr.append(numId)
                pPr.append(numPr)

                self._add_formatted_text(p, content)
                prev_style_type = style_type
            elif style_type == 'quote':
                try:
                    p = doc.add_paragraph(content, style='Quote')
                except KeyError:
                    # Quote style doesn't exist, use Normal with indent
                    p = doc.add_paragraph(content, style='Normal')
                    p.paragraph_format.left_indent = Inches(0.5)

                # Apply professional styling to blockquotes
                # Make text italic
                for run in p.runs:
                    run.italic = True

                # Add very light gray background
                self._add_paragraph_shading(p, 'F9F9F9')

                # Add left border (3pt blue)
                self._add_paragraph_border(p, ['left'], size='24', color='4A90E2', space='8')

                prev_style_type = style_type
            elif style_type == 'empty':
                # Track empty lines but don't add content
                prev_style_type = style_type
            elif style_type == 'normal' and content:
                p = doc.add_paragraph()
                self._add_formatted_text(p, content)
                prev_style_type = style_type

            i += 1

    def _parse_markdown_line(self, line: str):
        """Parse markdown line and return style, content, and indent level"""
        # Heading 1
        if line.startswith('# '):
            return 'Heading 1', line[2:].strip(), 0
        # Heading 2
        elif line.startswith('## '):
            return 'Heading 2', line[3:].strip(), 0
        # Heading 3
        elif line.startswith('### '):
            return 'Heading 3', line[4:].strip(), 0
        # Heading 4
        elif line.startswith('#### '):
            return 'Heading 4', line[5:].strip(), 0
        # Heading 5
        elif line.startswith('##### '):
            return 'Heading 5', line[6:].strip(), 0
        # Heading 6
        elif line.startswith('###### '):
            return 'Heading 6', line[7:].strip(), 0
        # Horizontal rule
        elif line.strip() == '---':
            return 'hr', '', 0
        # Table row
        elif '|' in line and line.strip().startswith('|'):
            return 'table', line.strip(), 0
        # Code block marker
        elif line.startswith('```'):
            return 'code_marker', line[3:].strip(), 0
        # Numbered list
        elif re.match(r'^\t*\d+\.\s+', line):
            level = self._get_indent_level(line)
            content = re.sub(r'^\t*\d+\.\s+', '', line).rstrip()
            return 'numbered_list', content, level
        # Bullet list
        elif re.match(r'^\t*[-*+]\s+', line):
            level = self._get_indent_level(line)
            content = re.sub(r'^\t*[-*+]\s+', '', line).rstrip()
            return 'list', content, level
        # Blockquote
        elif line.startswith('> '):
            return 'quote', line[2:].strip(), 0
        # Empty line
        elif line.strip() == '':
            return 'empty', '', 0
        # Normal paragraph
        else:
            return 'normal', line.strip(), 0

    def _get_indent_level(self, line: str) -> int:
        """Calculate indent level based on leading tabs"""
        leading_tabs = len(line) - len(line.lstrip('\t'))
        return leading_tabs

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code)"""
        parts = re.split(r'(\*\*.*?\*\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                # Add background shading to inline code
                self._add_run_shading(run, 'F5F5F5')
            else:
                code_parts = re.split(r'(`[^`]+`)', part)
                for code_part in code_parts:
                    if code_part.startswith('`') and code_part.endswith('`'):
                        run = paragraph.add_run(code_part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        # Add background shading to inline code
                        self._add_run_shading(run, 'F5F5F5')
                    elif code_part:
                        paragraph.add_run(code_part)
